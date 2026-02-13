"""
ClearBlueSky – RAG AI Knowledge (ChromaDB).
Chunk documents (.txt, .pdf, .md, .docx, .html, etc.), embed, and retrieve relevant excerpts for analysis.
"""

import os
import re
from pathlib import Path

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
RAG_DIR = Path(BASE_DIR) / "rag_store"
COLLECTION_NAME = "trading_books"
CHUNK_SIZE = 600
CHUNK_OVERLAP = 100

# Extensions we support with built-in loaders
TEXT_EXTENSIONS = {".txt", ".md", ".rst", ".markdown", ".csv", ".log", ".tex"}
HTML_EXTENSIONS = {".html", ".htm"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}
EPUB_EXTENSIONS = {".epub"}

SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | HTML_EXTENSIONS | PDF_EXTENSIONS | DOCX_EXTENSIONS | EPUB_EXTENSIONS

_chroma_ok = None


def _chroma_available():
    global _chroma_ok
    if _chroma_ok is None:
        try:
            import chromadb
            _chroma_ok = True
        except ImportError:
            _chroma_ok = False
    return _chroma_ok


def _chunk_text(text, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    """Split text into overlapping chunks."""
    chunks = []
    start = 0
    text = (text or "").strip().replace("\r\n", "\n")
    while start < len(text):
        end = start + size
        chunk = text[start:end]
        if not chunk.strip():
            start = end - overlap
            continue
        chunks.append(chunk.strip())
        start = end - overlap
    return chunks


def get_client():
    """Persistent ChromaDB client. Returns None if chromadb not installed."""
    if not _chroma_available():
        return None
    try:
        import chromadb
        RAG_DIR.mkdir(parents=True, exist_ok=True)
        return chromadb.PersistentClient(path=str(RAG_DIR))
    except Exception:
        return None


def _text_from_pdf(filepath):
    """Extract plain text from a PDF file. Returns str or None on error."""
    try:
        import fitz
        doc = fitz.open(filepath)
        parts = []
        for page in doc:
            parts.append(page.get_text() or "")
        doc.close()
        return "\n".join(parts).strip() or None
    except Exception:
        return None


def _text_from_html(filepath):
    """Extract plain text from HTML, stripping tags. Returns str or None."""
    try:
        raw = filepath.read_text(encoding="utf-8", errors="ignore")
        if not raw or not raw.strip():
            return None
        # Simple tag strip
        text = re.sub(r"<script[^>]*>[\s\S]*?</script>", " ", raw, flags=re.I)
        text = re.sub(r"<style[^>]*>[\s\S]*?</style>", " ", text, flags=re.I)
        text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"\s+", " ", text).strip()
        return text or None
    except Exception:
        return None


def _text_from_docx(filepath):
    """Extract plain text from DOCX. Returns str or None (needs python-docx)."""
    try:
        import docx
        doc = docx.Document(filepath)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells if cell.text.strip())
        return "\n".join(parts).strip() or None
    except ImportError:
        return None
    except Exception:
        return None


def _text_from_epub(filepath):
    """Extract plain text from EPUB. Returns str or None (needs ebooklib)."""
    try:
        import ebooklib
        from ebooklib import epub
        book = epub.read_epub(filepath)
        parts = []
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                try:
                    html = item.get_content().decode("utf-8", errors="ignore")
                    text = re.sub(r"<[^>]+>", " ", html)
                    text = re.sub(r"\s+", " ", text).strip()
                    if text:
                        parts.append(text)
                except Exception:
                    pass
        return "\n".join(parts).strip() or None
    except ImportError:
        return None
    except Exception:
        return None


def _extract_text(filepath, ext):
    """Extract text from a file based on extension. Returns str or None."""
    ext_lower = ext.lower()
    if ext_lower in TEXT_EXTENSIONS:
        try:
            return filepath.read_text(encoding="utf-8", errors="ignore").strip() or None
        except Exception:
            return None
    if ext_lower in PDF_EXTENSIONS:
        return _text_from_pdf(filepath)
    if ext_lower in HTML_EXTENSIONS:
        return _text_from_html(filepath)
    if ext_lower in DOCX_EXTENSIONS:
        return _text_from_docx(filepath)
    if ext_lower in EPUB_EXTENSIONS:
        return _text_from_epub(filepath)
    # Fallback: try reading as UTF-8 text (handles .notes, custom extensions, etc.)
    try:
        raw = filepath.read_bytes()
        text = raw.decode("utf-8", errors="replace")
        if not text or not text.strip():
            return None
        # Reject if mostly binary (low ratio of printable)
        printable = sum(1 for c in text[:2000] if c.isprintable() or c in "\n\r\t")
        if len(text[:2000]) > 0 and printable / len(text[:2000]) < 0.5:
            return None
        return text.strip() or None
    except Exception:
        return None


def build_index(knowledge_folder, progress_callback=None):
    """
    Index documents under knowledge_folder. Supports .txt, .pdf, .md, .html, .docx, .epub,
    and tries to read unknown extensions as text. Chunk and add to ChromaDB.
    progress_callback(msg) optional. Returns number of chunks added.
    """
    def progress(msg):
        if progress_callback:
            progress_callback(msg)
    client = get_client()
    if not client:
        progress("ChromaDB not installed; pip install chromadb")
        return 0
    knowledge_folder = Path(knowledge_folder) if knowledge_folder else None
    if not knowledge_folder or not knowledge_folder.is_dir():
        progress("No AI Knowledge folder")
        return 0
    try:
        try:
            client.delete_collection(COLLECTION_NAME)
        except Exception:
            pass
        collection = client.create_collection(name=COLLECTION_NAME, metadata={"description": "AI Knowledge excerpts"})
    except Exception as e:
        progress(f"ChromaDB error: {e}")
        return 0
    ids = []
    documents = []
    n = 0
    # Collect all candidate files (known extensions + any file for fallback)
    known_files = []
    fallback_files = []
    for fp in sorted(knowledge_folder.rglob("*")):
        if not fp.is_file():
            continue
        ext = fp.suffix
        if ext.lower() in SUPPORTED_EXTENSIONS:
            known_files.append(fp)
        elif ext and ext[0] == "." and len(ext) <= 8:
            fallback_files.append(fp)
    for fp in known_files + fallback_files:
        try:
            text = _extract_text(fp, fp.suffix)
            if not (text or "").strip():
                continue
            chunks = _chunk_text(text)
            for i, chunk in enumerate(chunks):
                doc_id = f"{fp.stem}_{n}_{i}"
                ids.append(doc_id)
                documents.append(chunk)
                n += 1
        except Exception:
            continue
    if not documents:
        progress("No document chunks found in AI Knowledge folder")
        return 0
    try:
        collection.add(ids=ids[:10000], documents=documents[:10000])  # cap for free tier
        progress(f"Indexed {len(documents)} chunks from {knowledge_folder}")
        return len(documents)
    except Exception as e:
        progress(f"Add error: {e}")
        return 0


def get_relevant_chunks(query, k=5):
    """
    Retrieve up to k most relevant chunks for query. Returns list of strings.
    If no index or chromadb, returns [].
    """
    client = get_client()
    if not client:
        return []
    try:
        collection = client.get_collection(name=COLLECTION_NAME)
        results = collection.query(query_texts=[query], n_results=min(k, 20))
        if not results or not results.get("documents"):
            return []
        return [doc for doc in (results["documents"][0] or []) if doc]
    except Exception:
        return []


def get_rag_context_for_scan(scan_type, analysis_summary="", k=5):
    """
    Build a query from scan_type and optional summary; return formatted string of chunks
    for inclusion in system prompt. If no chunks, returns "".
    """
    query = f"{scan_type} scan technical analysis pattern setup"
    if analysis_summary:
        query = f"{query} {analysis_summary[:200]}"
    chunks = get_relevant_chunks(query, k=k)
    if not chunks:
        return ""
    return "Relevant AI Knowledge excerpts (use for pattern/strategy context):\n" + "\n---\n".join(chunks[:k])
